#!/usr/bin/env python3
"""
RAG Engine - Core RAG functionality with tenant isolation
Handles document processing, embeddings, and querying
"""

import os
import json
import hashlib
import pickle
import asyncio
import aiofiles
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import requests
import sqlite3

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document
import pandas as pd
import re

class RAGEngine:
    """
    Core RAG engine with per-tenant isolation
    """
    
    def __init__(self, tenant_id: str, config: Optional[Dict] = None):
        """
        Initialize RAG engine for a specific tenant
        
        Args:
            tenant_id: Unique tenant identifier
            config: Optional configuration overrides
        """
        self.tenant_id = tenant_id
        
        # Default configuration
        self.config = {
            'model': 'llama3.2',
            'base_url': os.getenv('OLLAMA_HOST', 'http://localhost:11434'),
            'embedding_model': 'all-MiniLM-L6-v2',
            'chunk_size': 1500,
            'chunk_overlap': 200,
            'similarity_threshold': 0.15,
            'top_k_results': 5,
            'max_context_length': 4000,
            'temperature': 0.7,
            'cache_enabled': True,
            'cache_ttl': 3600  # 1 hour
        }
        
        # Override with provided config
        if config:
            self.config.update(config)
        
        # Initialize embedder (shared across all tenants for efficiency)
        print(f"🚀 Initializing RAG engine for tenant {tenant_id}...")
        self.embedder = SentenceTransformer(self.config['embedding_model'])
        
        # Tenant-specific storage paths
        self.data_dir = Path(f"./data/tenants/{tenant_id}")
        self.cache_dir = Path(f"./cache/tenants/{tenant_id}")
        self.data_dir.mkdir(parents=True, exist_ok=True)
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        # In-memory storage for this tenant
        self.chunks = []
        self.chunk_embeddings = None
        self.chunk_metadata = []
        
        # Query cache
        self.query_cache = {}
        
        # Load existing data if available
        self._load_tenant_data()
    
    def _load_tenant_data(self):
        """Load existing tenant data from disk"""
        embeddings_file = self.data_dir / "embeddings.pkl"
        metadata_file = self.data_dir / "metadata.json"
        chunks_file = self.data_dir / "chunks.json"
        
        if embeddings_file.exists():
            try:
                with open(embeddings_file, 'rb') as f:
                    self.chunk_embeddings = pickle.load(f)
                
                with open(metadata_file, 'r') as f:
                    self.chunk_metadata = json.load(f)
                
                with open(chunks_file, 'r') as f:
                    self.chunks = json.load(f)
                
                print(f"✅ Loaded {len(self.chunks)} existing chunks for tenant {self.tenant_id}")
            except Exception as e:
                print(f"⚠️ Error loading tenant data: {e}")
    
    def _save_tenant_data(self):
        """Save tenant data to disk"""
        try:
            embeddings_file = self.data_dir / "embeddings.pkl"
            metadata_file = self.data_dir / "metadata.json"
            chunks_file = self.data_dir / "chunks.json"
            
            with open(embeddings_file, 'wb') as f:
                pickle.dump(self.chunk_embeddings, f)
            
            with open(metadata_file, 'w') as f:
                json.dump(self.chunk_metadata, f)
            
            with open(chunks_file, 'w') as f:
                json.dump(self.chunks, f)
            
            print(f"✅ Saved {len(self.chunks)} chunks for tenant {self.tenant_id}")
        except Exception as e:
            print(f"❌ Error saving tenant data: {e}")
    
    async def extract_text(self, file_path: str) -> str:
        """
        Extract text from various file formats
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        filename = file_path.name.lower()
        
        try:
            if filename.endswith('.pdf'):
                text = ""
                # Try pdfplumber first (better for tables and complex layouts)
                try:
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n\n"
                except:
                    # Fallback to PyPDF2
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n\n"
                
                return text.strip()
            
            elif filename.endswith('.docx'):
                doc = Document(file_path)
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            paragraphs.append(row_text)
                
                return "\n\n".join(paragraphs)
            
            elif filename.endswith('.csv'):
                df = pd.read_csv(file_path)
                text_parts = [f"CSV Data: {filename}"]
                text_parts.append(f"Shape: {df.shape[0]} rows × {df.shape[1]} columns")
                text_parts.append(f"Columns: {', '.join(df.columns)}")
                text_parts.append("\nData:")
                text_parts.append(df.to_string(max_rows=100))
                return "\n".join(text_parts)
            
            elif filename.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path)
                text_parts = [f"Excel Data: {filename}"]
                text_parts.append(f"Shape: {df.shape[0]} rows × {df.shape[1]} columns")
                text_parts.append(f"Columns: {', '.join(df.columns)}")
                text_parts.append("\nData:")
                text_parts.append(df.to_string(max_rows=100))
                return "\n".join(text_parts)
            
            else:
                # Plain text files (txt, md, etc.)
                async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return await f.read()
        
        except Exception as e:
            print(f"❌ Error extracting text from {filename}: {e}")
            return f"[Error extracting content from {filename}]"
    
    def create_chunks(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Create intelligent chunks from text
        
        Args:
            text: Text content to chunk
            metadata: Metadata about the source document
            
        Returns:
            List of chunk dictionaries
        """
        chunks = []
        filename = metadata.get('filename', 'unknown')
        
        # Special handling for resumes/CVs
        if any(keyword in filename.lower() for keyword in ['resume', 'cv', 'curriculum']):
            # Keep the full document as first chunk for overview
            overview_chunk = {
                'text': text[:3000] if len(text) > 3000 else text,
                'metadata': {
                    **metadata,
                    'chunk_type': 'overview',
                    'chunk_index': 0
                }
            }
            chunks.append(overview_chunk)
        
        # Smart sentence-aware chunking
        sentences = re.split(r'(?<=[.!?])\s+', text)
        current_chunk = []
        current_size = 0
        chunk_index = len(chunks)
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            # If adding this sentence exceeds chunk size, save current chunk
            if current_size + sentence_size > self.config['chunk_size'] and current_chunk:
                chunk_text = ' '.join(current_chunk)
                chunks.append({
                    'text': chunk_text,
                    'metadata': {
                        **metadata,
                        'chunk_type': 'content',
                        'chunk_index': chunk_index
                    }
                })
                
                # Start new chunk with overlap
                overlap_sentences = []
                overlap_size = 0
                for sent in reversed(current_chunk):
                    if overlap_size + len(sent) <= self.config['chunk_overlap']:
                        overlap_sentences.insert(0, sent)
                        overlap_size += len(sent)
                    else:
                        break
                
                current_chunk = overlap_sentences + [sentence]
                current_size = overlap_size + sentence_size
                chunk_index += 1
            else:
                current_chunk.append(sentence)
                current_size += sentence_size
        
        # Add final chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            if len(chunk_text) > 50:  # Minimum chunk size
                chunks.append({
                    'text': chunk_text,
                    'metadata': {
                        **metadata,
                        'chunk_type': 'content',
                        'chunk_index': chunk_index
                    }
                })
        
        return chunks
    
    async def process_document(self, file_path: str, filename: str) -> int:
        """
        Process a document and add to tenant's knowledge base
        
        Args:
            file_path: Path to the document
            filename: Original filename
            
        Returns:
            Number of chunks created
        """
        # Extract text
        text = await self.extract_text(file_path)
        
        if not text or len(text) < 10:
            print(f"⚠️ No valid content extracted from {filename}")
            return 0
        
        # Create metadata
        metadata = {
            'filename': filename,
            'file_path': file_path,
            'processed_at': datetime.utcnow().isoformat(),
            'tenant_id': self.tenant_id
        }
        
        # Create chunks
        new_chunks = self.create_chunks(text, metadata)
        
        if not new_chunks:
            return 0
        
        # Add chunks to storage
        for chunk in new_chunks:
            self.chunks.append(chunk['text'])
            self.chunk_metadata.append(chunk['metadata'])
        
        # Generate embeddings for new chunks
        new_texts = [chunk['text'] for chunk in new_chunks]
        new_embeddings = self.embedder.encode(new_texts)
        
        # Update embeddings array
        if self.chunk_embeddings is None:
            self.chunk_embeddings = new_embeddings
        else:
            self.chunk_embeddings = np.vstack([self.chunk_embeddings, new_embeddings])
        
        # Save to disk
        self._save_tenant_data()
        
        # Clear query cache since new content was added
        self.query_cache.clear()
        
        print(f"✅ Processed {filename}: created {len(new_chunks)} chunks")
        return len(new_chunks)
    
    def search_similar_chunks(self, query: str, top_k: int = 5) -> List[Tuple[str, Dict, float]]:
        """
        Search for similar chunks using semantic similarity
        
        Args:
            query: Search query
            top_k: Number of results to return
            
        Returns:
            List of (chunk_text, metadata, similarity_score) tuples
        """
        if not self.chunks or self.chunk_embeddings is None:
            return []
        
        # Generate query embedding
        query_embedding = self.embedder.encode([query])
        
        # Calculate similarities
        similarities = cosine_similarity(query_embedding, self.chunk_embeddings)[0]
        
        # Get top-k indices
        top_indices = np.argsort(similarities)[-top_k:][::-1]
        
        # Filter by threshold and prepare results
        results = []
        for idx in top_indices:
            similarity = similarities[idx]
            if similarity >= self.config['similarity_threshold']:
                results.append((
                    self.chunks[idx],
                    self.chunk_metadata[idx],
                    float(similarity)
                ))
        
        return results
    
    async def query(self, query: str, top_k: int = 5, temperature: float = 0.7) -> Dict[str, Any]:
        """
        Query the knowledge base and generate a response
        
        Args:
            query: User query
            top_k: Number of context chunks to use
            temperature: LLM temperature
            
        Returns:
            Response dictionary with answer and metadata
        """
        # Check cache
        cache_key = hashlib.md5(f"{query}:{top_k}:{temperature}".encode()).hexdigest()
        if self.config['cache_enabled'] and cache_key in self.query_cache:
            cached = self.query_cache[cache_key]
            if (datetime.utcnow() - cached['timestamp']).seconds < self.config['cache_ttl']:
                return cached['response']
        
        # Search for relevant chunks
        similar_chunks = self.search_similar_chunks(query, top_k)
        
        if not similar_chunks:
            return {
                'answer': "I don't have enough information to answer your question. Please upload relevant documents first.",
                'sources': [],
                'confidence': 0.0,
                'tokens_used': 0
            }
        
        # Prepare context
        context_parts = []
        sources = []
        
        for chunk_text, metadata, similarity in similar_chunks:
            context_parts.append(chunk_text)
            sources.append({
                'filename': metadata.get('filename', 'Unknown'),
                'chunk_index': metadata.get('chunk_index', 0),
                'similarity': round(similarity, 3)
            })
        
        context = "\n\n---\n\n".join(context_parts)
        
        # Truncate context if too long
        if len(context) > self.config['max_context_length']:
            context = context[:self.config['max_context_length']]
        
        # Generate response using Ollama
        try:
            response = await self._generate_llm_response(query, context, temperature)
            
            # Calculate confidence based on similarity scores
            avg_similarity = np.mean([s[2] for s in similar_chunks])
            confidence = float(avg_similarity)
            
            result = {
                'answer': response['answer'],
                'sources': sources,
                'confidence': confidence,
                'tokens_used': response.get('tokens_used', 0)
            }
            
            # Cache the result
            if self.config['cache_enabled']:
                self.query_cache[cache_key] = {
                    'response': result,
                    'timestamp': datetime.utcnow()
                }
            
            return result
            
        except Exception as e:
            print(f"❌ Error generating response: {e}")
            return {
                'answer': "I encountered an error while processing your question. Please try again.",
                'sources': sources,
                'confidence': 0.0,
                'tokens_used': 0
            }
    
    async def _generate_llm_response(self, query: str, context: str, temperature: float) -> Dict[str, Any]:
        """
        Generate response using Ollama LLM
        
        Args:
            query: User query
            context: Context from documents
            temperature: LLM temperature
            
        Returns:
            Response dictionary
        """
        prompt = f"""You are a helpful AI assistant. Answer the following question based on the provided context.
If the answer cannot be found in the context, say so honestly.

Context:
{context}

Question: {query}

Answer:"""
        
        # Call Ollama API
        try:
            # Run blocking request in executor to make it async
            loop = asyncio.get_event_loop()
            response = await loop.run_in_executor(
                None,
                lambda: requests.post(
                    f"{self.config['base_url']}/api/generate",
                    json={
                        'model': self.config['model'],
                        'prompt': prompt,
                        'temperature': temperature,
                        'stream': False
                    },
                    timeout=30
                )
            )
            
            if response.status_code == 200:
                result = response.json()
                return {
                    'answer': result.get('response', 'No response generated'),
                    'tokens_used': len(prompt.split()) + len(result.get('response', '').split())
                }
            else:
                # Fallback to simple extraction if LLM fails
                return self._fallback_response(query, context)
                
        except Exception as e:
            print(f"⚠️ LLM request failed: {e}")
            return self._fallback_response(query, context)
    
    def _fallback_response(self, query: str, context: str) -> Dict[str, Any]:
        """
        Generate a simple fallback response without LLM
        
        Args:
            query: User query
            context: Context from documents
            
        Returns:
            Response dictionary
        """
        # Simple keyword-based extraction
        query_words = set(query.lower().split())
        sentences = context.split('.')
        
        relevant_sentences = []
        for sentence in sentences[:5]:  # Take first 5 sentences
            sentence_words = set(sentence.lower().split())
            if query_words & sentence_words:  # If there's overlap
                relevant_sentences.append(sentence.strip())
        
        if relevant_sentences:
            answer = '. '.join(relevant_sentences[:3]) + '.'
        else:
            answer = "Based on the documents: " + sentences[0].strip() if sentences else "No relevant information found."
        
        return {
            'answer': answer,
            'tokens_used': len(query.split()) + len(answer.split())
        }
    
    async def clear_tenant_data(self):
        """Clear all data for this tenant"""
        self.chunks = []
        self.chunk_embeddings = None
        self.chunk_metadata = []
        self.query_cache.clear()
        
        # Remove files
        for file in self.data_dir.glob("*"):
            file.unlink()
        
        print(f"✅ Cleared all data for tenant {self.tenant_id}")
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get statistics about the tenant's knowledge base"""
        return {
            'tenant_id': self.tenant_id,
            'total_chunks': len(self.chunks),
            'total_documents': len(set(m.get('filename') for m in self.chunk_metadata)),
            'cache_size': len(self.query_cache),
            'embedding_dimensions': self.chunk_embeddings.shape if self.chunk_embeddings is not None else None,
            'config': self.config
        }


# Utility functions for standalone testing
async def test_rag_engine():
    """Test the RAG engine"""
    engine = RAGEngine(tenant_id="test-tenant")
    
    # Test text extraction
    test_file = "test.txt"
    with open(test_file, 'w') as f:
        f.write("This is a test document. It contains information about RAG systems. RAG stands for Retrieval Augmented Generation.")
    
    # Process document
    chunks = await engine.process_document(test_file, "test.txt")
    print(f"Created {chunks} chunks")
    
    # Test query
    result = await engine.query("What does RAG stand for?")
    print(f"Answer: {result['answer']}")
    print(f"Confidence: {result['confidence']}")
    
    # Clean up
    os.remove(test_file)

if __name__ == "__main__":
    # Run test if executed directly
    asyncio.run(test_rag_engine())