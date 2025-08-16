#!/usr/bin/env python3
"""
FINAL RAG: Beautiful UI + Semantic Search + Smart Chunking
The complete package that actually works
"""

import gradio as gr
import requests
import json
import time
from pathlib import Path
from typing import List, Tuple, Dict
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import re
import hashlib
import PyPDF2
import pdfplumber
from docx import Document
import pandas as pd
import csv

class ProRAG:
    def __init__(self):
        # Initialize AI models
        self.model = "llama3.2"  # or mistral for commercial
        self.base_url = "http://localhost:11434"
        
        # Initialize semantic search
        print("🧠 Loading semantic search model...")
        self.embedder = SentenceTransformer('all-MiniLM-L6-v2')
        print("✅ Semantic search ready!")
        
        # Document storage
        self.documents = {}
        self.chunks = []
        self.chunk_embeddings = None
        self.chunk_metadata = []
        
    def smart_chunk(self, text: str, filename: str) -> List[Dict]:
        """Intelligent document chunking based on structure"""
        chunks = []
        
        # Clean text
        text = text.strip()
        
        # For resumes, keep the full context in first chunk
        if 'resume' in filename.lower() or 'cv' in filename.lower():
            # Add a full summary chunk first
            summary_chunk = text[:3000] if len(text) > 3000 else text
            chunks.append({
                'text': summary_chunk,
                'source': filename,
                'type': 'resume_full'
            })
        
        if filename.endswith('.md'):
            # Markdown: Split by headers for structure preservation
            sections = re.split(r'\n(?=#{1,6}\s+)', text)
            
            for section in sections:
                if len(section.strip()) > 50:
                    # Keep sections reasonable size
                    if len(section) > 1500:
                        # Split long sections by paragraphs
                        paragraphs = section.split('\n\n')
                        current = ""
                        for para in paragraphs:
                            if len(current) + len(para) < 1500:
                                current += para + "\n\n"
                            else:
                                if current:
                                    chunks.append({
                                        'text': current.strip(),
                                        'source': filename,
                                        'type': 'md_section'
                                    })
                                current = para
                        if current:
                            chunks.append({
                                'text': current.strip(),
                                'source': filename,
                                'type': 'md_section'
                            })
                    else:
                        chunks.append({
                            'text': section.strip(),
                            'source': filename,
                            'type': 'md_section'
                        })
        
        elif filename.endswith(('.txt', '.text')):
            # Text: Split by paragraphs for natural breaks
            paragraphs = text.split('\n\n')
            current_chunk = ""
            
            for para in paragraphs:
                if len(current_chunk) + len(para) < 1200:
                    current_chunk += para + "\n\n"
                else:
                    if current_chunk:
                        chunks.append({
                            'text': current_chunk.strip(),
                            'source': filename,
                            'type': 'paragraph'
                        })
                    current_chunk = para
            
            if current_chunk:
                chunks.append({
                    'text': current_chunk.strip(),
                    'source': filename,
                    'type': 'paragraph'
                })
        
        else:
            # Generic: Split by size with overlap
            words = text.split()
            chunk_size = 200  # words
            overlap = 30  # words
            
            for i in range(0, len(words), chunk_size - overlap):
                chunk_text = ' '.join(words[i:i + chunk_size])
                if chunk_text:
                    chunks.append({
                        'text': chunk_text,
                        'source': filename,
                        'type': 'generic'
                    })
        
        return chunks if chunks else [{'text': text[:3000], 'source': filename, 'type': 'fallback'}]
    
    def extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text = ""
        
        try:
            # Try pdfplumber first (better for tables and formatting)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
        except:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text += page.extract_text() + "\n\n"
            except Exception as e:
                print(f"Error extracting PDF: {e}")
                text = "[PDF extraction failed]"
        
        return text.strip()
    
    def extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n\n".join(text)
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
            return "[DOCX extraction failed]"
    
    def extract_csv_text(self, file_path: str) -> str:
        """Extract and format CSV data"""
        try:
            # Read CSV with pandas
            df = pd.read_csv(file_path)
            
            # Create a text representation
            text_parts = []
            
            # Add column names as header
            text_parts.append(f"Columns: {', '.join(df.columns)}")
            text_parts.append(f"Total rows: {len(df)}")
            text_parts.append("")
            
            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_parts.append("Summary Statistics:")
                for col in numeric_cols:
                    text_parts.append(f"  {col}: Mean={df[col].mean():.2f}, Min={df[col].min()}, Max={df[col].max()}")
                text_parts.append("")
            
            # Convert data to readable format
            text_parts.append("Data:")
            
            # For large CSVs, sample the data
            if len(df) > 100:
                # Take first 50 and last 50 rows
                sample_df = pd.concat([df.head(50), df.tail(50)])
                text_parts.append(f"(Showing first 50 and last 50 of {len(df)} rows)")
            else:
                sample_df = df
            
            # Convert to string with nice formatting
            for idx, row in sample_df.iterrows():
                row_text = f"Row {idx}: " + " | ".join([f"{col}={val}" for col, val in row.items()])
                text_parts.append(row_text)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            print(f"Error extracting CSV: {e}")
            return "[CSV extraction failed]"
    
    def extract_excel_text(self, file_path: str) -> str:
        """Extract and format Excel data"""
        try:
            # Special handling for Google Sheets exports
            # These are often HTML files disguised as XLS
            
            # First, try to detect if it's actually HTML
            with open(file_path, 'rb') as f:
                first_bytes = f.read(100)
                if b'<html' in first_bytes.lower() or b'<!doctype' in first_bytes.lower():
                    print(f"📋 Detected HTML table format (Google Sheets export)")
                    # Read as HTML table
                    try:
                        dfs = pd.read_html(file_path)
                        if dfs:
                            df = dfs[0]  # Get first table
                            text_parts = []
                            text_parts.append(f"Google Sheets Export: {Path(file_path).name}")
                            text_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
                            text_parts.append(f"Columns: {', '.join(map(str, df.columns))}")
                            text_parts.append("\nData preview:")
                            text_parts.append(df.head(20).to_string())
                            return "\n".join(text_parts)
                    except Exception as e:
                        print(f"Error reading HTML table: {e}")
            
            # Try normal Excel reading
            try:
                # Try openpyxl for xlsx
                excel_file = pd.ExcelFile(file_path, engine='openpyxl')
            except:
                try:
                    # Try xlrd for older xls files
                    excel_file = pd.ExcelFile(file_path, engine='xlrd')
                except:
                    try:
                        # Try reading directly with pandas (auto-detect)
                        df = pd.read_excel(file_path, engine=None)
                        text_parts = []
                        text_parts.append(f"Spreadsheet: {Path(file_path).name}")
                        text_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
                        text_parts.append(f"Columns: {', '.join(map(str, df.columns))}")
                        text_parts.append("\nData preview:")
                        text_parts.append(df.to_string())
                        return "\n".join(text_parts)
                    except:
                        # Let pandas choose the engine
                        excel_file = pd.ExcelFile(file_path)
            
            text_parts = []
            
            text_parts.append(f"Excel file: {Path(file_path).name}")
            text_parts.append(f"Sheets: {len(excel_file.sheet_names)}")
            text_parts.append("")
            
            # Process each sheet
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    text_parts.append(f"=== Sheet: {sheet_name} ===")
                    text_parts.append(f"Columns: {', '.join(map(str, df.columns))}")
                    text_parts.append(f"Rows: {len(df)}")
                    
                    # Add summary for numeric columns
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        text_parts.append("Summary:")
                        for col in numeric_cols:
                            if df[col].notna().any():  # Check if column has any non-null values
                                mean_val = df[col].mean()
                                text_parts.append(f"  {col}: Mean={mean_val:.2f}, Min={df[col].min()}, Max={df[col].max()}")
                    
                    # Sample data from each sheet
                    if len(df) > 50:
                        sample_df = pd.concat([df.head(25), df.tail(25)])
                        text_parts.append(f"\nData (showing 50 of {len(df)} rows):")
                    else:
                        sample_df = df
                        text_parts.append("\nData:")
                    
                    for idx, row in sample_df.iterrows():
                        row_text = " | ".join([f"{col}={val}" for col, val in row.items() if pd.notna(val)])
                        text_parts.append(f"  Row {idx}: {row_text}")
                    
                    text_parts.append("")
                    
                except Exception as e:
                    text_parts.append(f"Error reading sheet {sheet_name}: {e}")
                    text_parts.append("")
            
            if len(text_parts) > 3:  # Has some content
                return "\n".join(text_parts)
            else:
                return f"Excel file detected but could not extract content. File: {Path(file_path).name}"
            
        except Exception as e:
            print(f"Error extracting Excel: {e}")
            # Try a simple fallback
            try:
                # Just try to read first sheet with basic settings
                df = pd.read_excel(file_path, sheet_name=0, engine=None)
                text = f"Excel file with {len(df)} rows and {len(df.columns)} columns\n"
                text += f"Columns: {', '.join(map(str, df.columns))}\n"
                text += f"First few rows:\n{df.head(10).to_string()}"
                return text
            except:
                return f"[Excel extraction failed: {e}]"
    
    def load_documents(self, files):
        """Load documents with smart chunking and embeddings"""
        if not files:
            return gr.update(value="📎 No files selected"), gr.update(visible=False), "No documents"
        
        self.documents = {}
        self.chunks = []
        self.chunk_metadata = []
        loaded = []
        
        for file in files:
            try:
                filename = Path(file.name).name
                
                # Extract content based on file type
                if filename.lower().endswith('.pdf'):
                    print(f"📑 Extracting PDF: {filename}")
                    content = self.extract_pdf_text(file.name)
                elif filename.lower().endswith('.docx'):
                    print(f"📄 Extracting DOCX: {filename}")
                    content = self.extract_docx_text(file.name)
                elif filename.lower().endswith('.csv'):
                    print(f"📊 Extracting CSV: {filename}")
                    content = self.extract_csv_text(file.name)
                elif filename.lower().endswith(('.xlsx', '.xls')):
                    print(f"📈 Extracting Excel: {filename}")
                    content = self.extract_excel_text(file.name)
                    if not content or "failed" in content.lower():
                        print(f"⚠️ Excel extraction had issues, content length: {len(content) if content else 0}")
                else:
                    # Text-based files (txt, md, etc.)
                    with open(file.name, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                
                if content:
                    self.documents[filename] = content
                
                # Smart chunking
                file_chunks = self.smart_chunk(content, filename)
                
                # Add to chunk list
                for chunk in file_chunks:
                    self.chunks.append(chunk['text'])
                    self.chunk_metadata.append({
                        'source': chunk['source'],
                        'type': chunk['type'],
                        'length': len(chunk['text'])
                    })
                
                loaded.append(filename)
                print(f"📄 Loaded {filename}: {len(file_chunks)} chunks")
                
            except Exception as e:
                print(f"Error: {e}")
        
        if self.chunks:
            # Create embeddings for semantic search
            print(f"🔍 Creating embeddings for {len(self.chunks)} chunks...")
            self.chunk_embeddings = self.embedder.encode(self.chunks)
            print("✅ Embeddings created!")
            
            return (
                gr.update(value=f"✅ {len(loaded)} files ready"),
                gr.update(visible=True, value=f"📚 Loaded: {', '.join(loaded)} ({len(self.chunks)} searchable chunks)"),
                f"{len(self.chunks)} chunks indexed"
            )
        
        return gr.update(value="❌ Failed"), gr.update(visible=False), "No chunks"
    
    def semantic_search(self, query: str, top_k: int = 5) -> List[Dict]:
        """Find relevant chunks using semantic similarity"""
        if self.chunk_embeddings is None or len(self.chunks) == 0:
            return []
        
        # Encode query
        query_embedding = self.embedder.encode([query])
        
        # Calculate similarities
        similarities = cosine_similarity(query_embedding, self.chunk_embeddings)[0]
        
        # Get top k indices
        top_indices = np.argsort(similarities)[-top_k:][::-1]
        
        # Return chunks with scores
        results = []
        for idx in top_indices:
            if similarities[idx] > 0.1:  # Lowered threshold for better recall
                results.append({
                    'text': self.chunks[idx],
                    'score': float(similarities[idx]),
                    'metadata': self.chunk_metadata[idx]
                })
        
        # If no results with threshold, return top result anyway
        if not results and len(top_indices) > 0:
            idx = top_indices[0]
            results.append({
                'text': self.chunks[idx],
                'score': float(similarities[idx]),
                'metadata': self.chunk_metadata[idx]
            })
        
        return results
    
    def stream_response(self, message: str, history: List[Tuple[str, str]]):
        """Stream response using semantic search"""
        if not message.strip():
            return history
        
        history.append([message, None])
        
        # Use SEMANTIC SEARCH to find relevant chunks
        if self.chunks:
            print(f"🔍 Searching for: {message}")
            relevant_chunks = self.semantic_search(message, top_k=5)
            
            if relevant_chunks:
                # Build context from semantically similar chunks
                context_parts = []
                sources = set()
                
                for chunk in relevant_chunks:
                    context_parts.append(chunk['text'])
                    sources.add(chunk['metadata']['source'])
                    print(f"  Found relevant chunk (score: {chunk['score']:.2f})")
                
                context = "\n\n".join(context_parts)
                prompt = f"""Based on the following context, answer the question.

Context:
{context[:4000]}

Question: {message}

Answer:"""
                
                # Add source attribution
                source_text = f"\n\n📄 Sources: {', '.join(sources)}"
            else:
                # No relevant chunks found
                prompt = message
                source_text = "\n\n⚠️ No relevant sections found in documents"
        else:
            prompt = message
            source_text = ""
        
        # Query Ollama
        try:
            response = requests.post(
                f"{self.base_url}/api/generate",
                json={"model": self.model, "prompt": prompt, "stream": False}
            )
            
            if response.status_code == 200:
                full_response = response.json().get('response', 'No response')
                full_response += source_text
                
                # Typing effect
                history[-1][1] = ""
                for i in range(0, len(full_response), 5):
                    history[-1][1] = full_response[:i+5]
                    yield history
                    time.sleep(0.01)
            else:
                history[-1][1] = "❌ Error: Ollama not running"
                yield history
                
        except Exception as e:
            history[-1][1] = f"❌ Error: {str(e)}"
            yield history

def create_ui():
    """Create the beautiful UI with proper RAG backend"""
    
    rag = ProRAG()
    
    # Modern CSS (same beautiful design)
    custom_css = """
    :root {
        --primary: #10B981;
        --primary-hover: #059669;
        --bg-primary: #0F172A;
        --bg-secondary: #1E293B;
        --bg-tertiary: #334155;
        --text-primary: #F1F5F9;
        --text-secondary: #94A3B8;
        --border: #334155;
        --accent: #6366F1;
    }
    
    .gradio-container {
        background: linear-gradient(180deg, #0F172A 0%, #1E293B 100%) !important;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
        color: var(--text-primary) !important;
    }
    
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 1rem;
        margin-bottom: 2rem;
        box-shadow: 0 20px 40px rgba(0,0,0,0.3);
        animation: fadeIn 0.5s ease-in;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(-10px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    .message {
        animation: slideIn 0.3s ease-out;
    }
    
    @keyframes slideIn {
        from { opacity: 0; transform: translateX(-10px); }
        to { opacity: 1; transform: translateX(0); }
    }
    
    .primary-btn {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-hover) 100%) !important;
        color: white !important;
        border: none !important;
        padding: 0.75rem 2rem !important;
        border-radius: 0.75rem !important;
        font-weight: 600 !important;
        cursor: pointer !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 6px rgba(16, 185, 129, 0.2) !important;
    }
    
    .primary-btn:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 12px rgba(16, 185, 129, 0.3) !important;
    }
    
    .file-upload {
        background: var(--bg-tertiary) !important;
        border: 2px dashed var(--border) !important;
        border-radius: 1rem !important;
        padding: 2rem !important;
        transition: all 0.3s ease !important;
    }
    
    .file-upload:hover {
        border-color: var(--primary) !important;
        background: rgba(16, 185, 129, 0.05) !important;
    }
    """
    
    with gr.Blocks(theme=gr.themes.Base(), css=custom_css, title="AI Document Assistant") as app:
        
        # Header
        gr.HTML("""
        <div class="main-header">
            <h1 style="color: white; margin: 0; font-size: 2.5rem; font-weight: 700;">
                ✨ AI Document Assistant
            </h1>
            <p style="color: rgba(255,255,255,0.9); margin-top: 0.5rem; font-size: 1.1rem;">
                Semantic Search + Smart Chunking + Beautiful UI
            </p>
            <div style="margin-top: 1rem;">
                <span style="background: rgba(16, 185, 129, 0.2); color: #10B981; padding: 0.25rem 0.75rem; border-radius: 9999px; font-size: 0.875rem;">
                    ● Semantic Search Active
                </span>
            </div>
        </div>
        """)
        
        with gr.Row():
            # Main Chat
            with gr.Column(scale=3):
                chatbot = gr.Chatbot(
                    height=600,
                    show_label=False,
                    avatar_images=["🧑‍💼", "🤖"]
                )
                
                with gr.Row():
                    msg = gr.Textbox(
                        placeholder="Ask anything... (I'll search semantically through your documents)",
                        show_label=False,
                        scale=4,
                        lines=1
                    )
                    send_btn = gr.Button("Send →", elem_classes="primary-btn", scale=1)
                
                with gr.Row():
                    clear_btn = gr.Button("🗑️ Clear")
                    
                    gr.Examples(
                        examples=[
                            "Summarize my experience",
                            "What are my key skills?",
                            "What projects have I worked on?",
                            "What is my education background?",
                            "Create a brief bio from this"
                        ],
                        inputs=msg
                    )
            
            # Right Sidebar
            with gr.Column(scale=1):
                # Document Upload
                with gr.Group():
                    gr.Markdown("### 📚 Documents")
                    
                    file_upload = gr.File(
                        label="Upload Files",
                        file_count="multiple",
                        file_types=[".txt", ".pdf", ".md", ".docx", ".csv", ".xlsx", ".xls"],
                        elem_classes="file-upload"
                    )
                    
                    upload_btn = gr.Button("📤 Process Documents", elem_classes="primary-btn")
                    
                    upload_status = gr.Textbox(
                        value="📎 No documents",
                        show_label=False,
                        interactive=False
                    )
                    
                    file_info = gr.Markdown(visible=False)
                    
                    chunk_info = gr.Textbox(
                        value="No chunks indexed",
                        label="Search Index",
                        interactive=False
                    )
                
                # Info
                with gr.Group():
                    gr.Markdown("""
                    ### 🧠 How It Works
                    
                    1. **Smart Chunking**: Documents split intelligently
                    2. **Semantic Search**: Finds meaning, not keywords
                    3. **Context Aware**: Uses relevant sections only
                    
                    ### ✅ Features Active
                    
                    - Semantic embeddings
                    - Smart document parsing
                    - Context ranking
                    - Source attribution
                    """)
        
        # Event handlers
        def process_message(msg, history):
            if not msg.strip():
                return "", history
            yield "", history
            for updated_history in rag.stream_response(msg, history):
                yield "", updated_history
        
        msg.submit(process_message, [msg, chatbot], [msg, chatbot])
        send_btn.click(process_message, [msg, chatbot], [msg, chatbot])
        clear_btn.click(lambda: ([], ""), outputs=[chatbot, msg])
        
        upload_btn.click(
            rag.load_documents,
            inputs=[file_upload],
            outputs=[upload_status, file_info, chunk_info]
        )
    
    return app

if __name__ == "__main__":
    print("\n" + "="*60)
    print("🚀 FINAL RAG: Complete Package")
    print("="*60)
    print("✅ Semantic Search - Understands meaning")
    print("✅ Smart Chunking - Preserves structure")
    print("✅ Beautiful UI - Professional design")
    print("✅ Source Attribution - Shows where answers come from")
    print("="*60 + "\n")
    
    app = create_ui()
    app.launch(server_name="0.0.0.0", server_port=7860)