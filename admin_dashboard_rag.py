#!/usr/bin/env python3
"""
RAG Admin Dashboard with Configuration & Debugging
Complete control panel for managing and troubleshooting your RAG system
"""

import gradio as gr
import requests
import json
import time
from pathlib import Path
from typing import List, Tuple, Dict
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import re
import hashlib
import PyPDF2
import pdfplumber
from docx import Document
import pandas as pd
import csv
import pickle
import yaml
from datetime import datetime

class AdminRAG:
    def __init__(self):
        # Configuration (adjustable via dashboard)
        self.config = {
            'model': 'llama3.2',
            'base_url': 'http://localhost:11434',
            'similarity_threshold': 0.15,
            'chunk_size': 1500,
            'chunk_overlap': 200,
            'top_k_results': 5,
            'min_chunk_size': 50,
            'max_context_length': 4000,
            'enable_resume_mode': True,
            'debug_mode': True
        }
        
        # Initialize AI models
        print("🧠 Loading semantic search model...")
        self.embedder = SentenceTransformer('all-MiniLM-L6-v2')
        print("✅ Semantic search ready!")
        
        # Document storage
        self.documents = {}
        self.chunks = []
        self.chunk_embeddings = None
        self.chunk_metadata = []
        
        # Debug information
        self.last_query_debug = {}
        self.processing_stats = {
            'total_chunks': 0,
            'avg_chunk_size': 0,
            'total_documents': 0,
            'last_processing_time': 0
        }
        
    def save_config(self, filename='rag_config.yaml'):
        """Save current configuration to file"""
        with open(filename, 'w') as f:
            yaml.dump(self.config, f)
        return f"✅ Configuration saved to {filename}"
    
    def load_config(self, filename='rag_config.yaml'):
        """Load configuration from file"""
        try:
            with open(filename, 'r') as f:
                self.config = yaml.load(f, Loader=yaml.FullLoader)
            return f"✅ Configuration loaded from {filename}"
        except Exception as e:
            return f"❌ Error loading config: {e}"
    
    def update_config(self, **kwargs):
        """Update configuration values"""
        for key, value in kwargs.items():
            if key in self.config:
                self.config[key] = value
        return self.config
    
    def smart_chunk(self, text: str, filename: str) -> List[Dict]:
        """Intelligent document chunking with configurable parameters"""
        chunks = []
        chunk_size = self.config['chunk_size']
        overlap = self.config['chunk_overlap']
        
        # Clean text
        text = text.strip()
        
        # Resume mode - keep full context in first chunk
        if self.config['enable_resume_mode'] and ('resume' in filename.lower() or 'cv' in filename.lower()):
            summary_chunk = text[:chunk_size * 2] if len(text) > chunk_size * 2 else text
            chunks.append({
                'text': summary_chunk,
                'source': filename,
                'type': 'resume_full',
                'size': len(summary_chunk)
            })
        
        if filename.endswith('.md'):
            # Markdown: Split by headers
            sections = re.split(r'\n(?=#{1,6}\s+)', text)
            
            for section in sections:
                if len(section.strip()) > self.config['min_chunk_size']:
                    if len(section) > chunk_size:
                        # Split long sections
                        words = section.split()
                        for i in range(0, len(words), chunk_size - overlap):
                            chunk_text = ' '.join(words[i:i + chunk_size])
                            if chunk_text:
                                chunks.append({
                                    'text': chunk_text,
                                    'source': filename,
                                    'type': 'md_section',
                                    'size': len(chunk_text)
                                })
                    else:
                        chunks.append({
                            'text': section.strip(),
                            'source': filename,
                            'type': 'md_section',
                            'size': len(section.strip())
                        })
        else:
            # Generic chunking with configurable size
            words = text.split()
            word_chunk_size = chunk_size // 5  # Approximate words
            word_overlap = overlap // 5
            
            for i in range(0, len(words), word_chunk_size - word_overlap):
                chunk_text = ' '.join(words[i:i + word_chunk_size])
                if len(chunk_text) > self.config['min_chunk_size']:
                    chunks.append({
                        'text': chunk_text,
                        'source': filename,
                        'type': 'generic',
                        'size': len(chunk_text)
                    })
        
        return chunks if chunks else [{'text': text[:chunk_size], 'source': filename, 'type': 'fallback', 'size': len(text[:chunk_size])}]
    
    def extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
        except:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text += page.extract_text() + "\n\n"
            except Exception as e:
                print(f"Error extracting PDF: {e}")
                text = "[PDF extraction failed]"
        return text.strip()
    
    def extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            return "\n\n".join(text)
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
            return "[DOCX extraction failed]"
    
    def load_documents(self, files):
        """Load documents with detailed statistics"""
        if not files:
            return gr.update(value="📎 No files selected"), "", {}
        
        start_time = time.time()
        self.documents = {}
        self.chunks = []
        self.chunk_metadata = []
        loaded = []
        stats = []
        
        for file in files:
            try:
                filename = Path(file.name).name
                
                # Extract content based on file type
                if filename.lower().endswith('.pdf'):
                    content = self.extract_pdf_text(file.name)
                elif filename.lower().endswith('.docx'):
                    content = self.extract_docx_text(file.name)
                else:
                    with open(file.name, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                
                if content:
                    self.documents[filename] = content
                
                # Smart chunking
                file_chunks = self.smart_chunk(content, filename)
                
                # Add to chunk list
                for chunk in file_chunks:
                    self.chunks.append(chunk['text'])
                    self.chunk_metadata.append({
                        'source': chunk['source'],
                        'type': chunk['type'],
                        'size': chunk['size']
                    })
                
                loaded.append(filename)
                stats.append(f"📄 {filename}: {len(file_chunks)} chunks, {len(content)} chars")
                
            except Exception as e:
                stats.append(f"❌ Error loading {filename}: {e}")
        
        if self.chunks:
            # Create embeddings
            print(f"🔍 Creating embeddings for {len(self.chunks)} chunks...")
            self.chunk_embeddings = self.embedder.encode(self.chunks)
            print("✅ Embeddings created!")
            
            # Update statistics
            self.processing_stats['total_chunks'] = len(self.chunks)
            self.processing_stats['avg_chunk_size'] = sum(m['size'] for m in self.chunk_metadata) / len(self.chunks)
            self.processing_stats['total_documents'] = len(loaded)
            self.processing_stats['last_processing_time'] = time.time() - start_time
            
            return (
                gr.update(value=f"✅ {len(loaded)} files loaded"),
                "\n".join(stats),
                self.processing_stats
            )
        
        return gr.update(value="❌ No content extracted"), "\n".join(stats), {}
    
    def semantic_search_with_debug(self, query: str, top_k: int = None) -> Tuple[List[Dict], Dict]:
        """Semantic search with detailed debug information"""
        if self.chunk_embeddings is None or len(self.chunks) == 0:
            return [], {"error": "No documents loaded"}
        
        if top_k is None:
            top_k = self.config['top_k_results']
        
        # Encode query
        query_embedding = self.embedder.encode([query])
        
        # Calculate similarities
        similarities = cosine_similarity(query_embedding, self.chunk_embeddings)[0]
        
        # Get top k indices
        top_indices = np.argsort(similarities)[-top_k:][::-1]
        
        # Debug information
        debug_info = {
            'query': query,
            'query_embedding_norm': float(np.linalg.norm(query_embedding)),
            'total_chunks_searched': len(self.chunks),
            'threshold_used': self.config['similarity_threshold'],
            'top_k_requested': top_k,
            'similarity_scores': []
        }
        
        # Return chunks with scores
        results = []
        for idx in top_indices:
            score = float(similarities[idx])
            debug_info['similarity_scores'].append({
                'chunk_index': int(idx),
                'score': score,
                'above_threshold': score > self.config['similarity_threshold'],
                'chunk_type': self.chunk_metadata[idx]['type'],
                'chunk_size': self.chunk_metadata[idx]['size']
            })
            
            if score > self.config['similarity_threshold']:
                results.append({
                    'text': self.chunks[idx],
                    'score': score,
                    'metadata': self.chunk_metadata[idx]
                })
        
        # If no results meet threshold, return top result anyway
        if not results and len(top_indices) > 0:
            idx = top_indices[0]
            results.append({
                'text': self.chunks[idx],
                'score': float(similarities[idx]),
                'metadata': self.chunk_metadata[idx]
            })
            debug_info['fallback_used'] = True
        
        self.last_query_debug = debug_info
        return results, debug_info
    
    def test_query(self, query: str) -> str:
        """Test a query and return formatted debug information"""
        if not query:
            return "Please enter a query to test"
        
        results, debug_info = self.semantic_search_with_debug(query)
        
        output = []
        output.append(f"🔍 Query: '{query}'")
        output.append(f"📊 Threshold: {self.config['similarity_threshold']}")
        output.append(f"🎯 Top K: {self.config['top_k_results']}")
        output.append("-" * 50)
        
        if 'error' in debug_info:
            output.append(f"❌ Error: {debug_info['error']}")
        else:
            output.append(f"📚 Searched {debug_info['total_chunks_searched']} chunks")
            output.append(f"📏 Query embedding norm: {debug_info['query_embedding_norm']:.3f}")
            
            if debug_info.get('fallback_used'):
                output.append("⚠️ No results above threshold - using best match")
            
            output.append("\n🎯 Top Matches:")
            for i, score_info in enumerate(debug_info['similarity_scores'][:5], 1):
                status = "✅" if score_info['above_threshold'] else "❌"
                output.append(
                    f"  {i}. {status} Score: {score_info['score']:.4f} | "
                    f"Type: {score_info['chunk_type']} | "
                    f"Size: {score_info['chunk_size']} chars"
                )
            
            if results:
                output.append(f"\n📋 Returned {len(results)} chunks for context")
                output.append("\n🔍 Preview of best match:")
                output.append(f"Score: {results[0]['score']:.4f}")
                output.append(f"Text: {results[0]['text'][:200]}...")
        
        return "\n".join(output)
    
    def get_system_status(self) -> str:
        """Get current system status and configuration"""
        status = []
        status.append("🎛️ SYSTEM STATUS")
        status.append("=" * 50)
        
        # Model status
        try:
            response = requests.get(f"{self.config['base_url']}/api/tags", timeout=2)
            if response.status_code == 200:
                status.append(f"✅ Ollama: Online")
                models = response.json().get('models', [])
                if models:
                    status.append(f"📦 Available models: {', '.join([m['name'] for m in models[:3]])}")
            else:
                status.append(f"⚠️ Ollama: Offline")
        except:
            status.append(f"❌ Ollama: Not reachable")
        
        status.append("")
        status.append("📊 STATISTICS")
        status.append("-" * 30)
        status.append(f"📚 Documents: {self.processing_stats['total_documents']}")
        status.append(f"📄 Chunks: {self.processing_stats['total_chunks']}")
        status.append(f"📏 Avg chunk size: {self.processing_stats['avg_chunk_size']:.0f} chars")
        status.append(f"⏱️ Last processing: {self.processing_stats['last_processing_time']:.2f}s")
        
        status.append("")
        status.append("⚙️ CONFIGURATION")
        status.append("-" * 30)
        for key, value in self.config.items():
            status.append(f"{key}: {value}")
        
        return "\n".join(status)
    
    def export_chunks(self) -> str:
        """Export chunks and embeddings for analysis"""
        if not self.chunks:
            return "No chunks to export"
        
        export_data = {
            'chunks': self.chunks,
            'metadata': self.chunk_metadata,
            'config': self.config,
            'timestamp': datetime.now().isoformat()
        }
        
        filename = f"rag_chunks_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        return f"✅ Exported {len(self.chunks)} chunks to {filename}"
    
    def stream_response(self, message: str, history: List[Tuple[str, str]]):
        """Stream response using configurable semantic search"""
        if not message.strip():
            return history
        
        history.append([message, None])
        
        # Use semantic search with current config
        if self.chunks:
            results, debug_info = self.semantic_search_with_debug(message)
            
            if results:
                context_parts = []
                sources = set()
                
                for chunk in results:
                    context_parts.append(chunk['text'])
                    sources.add(chunk['metadata']['source'])
                
                context = "\n\n".join(context_parts)[:self.config['max_context_length']]
                prompt = f"""Based on the following context, answer the question.

Context:
{context}

Question: {message}

Answer:"""
                
                source_text = f"\n\n📄 Sources: {', '.join(sources)}"
                
                if self.config['debug_mode']:
                    source_text += f"\n🔍 Debug: Found {len(results)} chunks, best score: {results[0]['score']:.3f}"
            else:
                prompt = message
                source_text = "\n\n⚠️ No relevant sections found"
        else:
            prompt = message
            source_text = ""
        
        # Query Ollama
        try:
            response = requests.post(
                f"{self.config['base_url']}/api/generate",
                json={"model": self.config['model'], "prompt": prompt, "stream": False}
            )
            
            if response.status_code == 200:
                full_response = response.json().get('response', 'No response')
                full_response += source_text
                
                history[-1][1] = ""
                for i in range(0, len(full_response), 5):
                    history[-1][1] = full_response[:i+5]
                    yield history
                    time.sleep(0.01)
            else:
                history[-1][1] = "❌ Error: Model not responding"
                yield history
                
        except Exception as e:
            history[-1][1] = f"❌ Error: {str(e)}"
            yield history

def create_admin_dashboard():
    """Create the admin dashboard with configuration controls"""
    
    rag = AdminRAG()
    
    custom_css = """
    :root {
        --primary: #6366F1;
        --success: #10B981;
        --warning: #F59E0B;
        --danger: #EF4444;
        --dark: #0F172A;
        --dark-secondary: #1E293B;
        --text: #F1F5F9;
    }
    
    .gradio-container {
        background: linear-gradient(180deg, #0F172A 0%, #1E293B 100%) !important;
        color: var(--text) !important;
    }
    
    .admin-header {
        background: linear-gradient(135deg, #6366F1 0%, #8B5CF6 100%);
        padding: 2rem;
        border-radius: 1rem;
        margin-bottom: 2rem;
        box-shadow: 0 20px 40px rgba(0,0,0,0.3);
    }
    
    .config-card {
        background: var(--dark-secondary);
        padding: 1.5rem;
        border-radius: 0.75rem;
        border: 1px solid rgba(255,255,255,0.1);
        margin-bottom: 1rem;
    }
    
    .status-badge {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 9999px;
        font-size: 0.875rem;
        font-weight: 500;
    }
    
    .debug-output {
        background: #000;
        color: #0F0;
        font-family: monospace;
        padding: 1rem;
        border-radius: 0.5rem;
        border: 1px solid #0F0;
    }
    """
    
    with gr.Blocks(theme=gr.themes.Base(), css=custom_css, title="RAG Admin Dashboard") as app:
        
        # Header
        gr.HTML("""
        <div class="admin-header">
            <h1 style="color: white; margin: 0; font-size: 2.5rem; font-weight: 700;">
                🎛️ RAG Admin Dashboard
            </h1>
            <p style="color: rgba(255,255,255,0.9); margin-top: 0.5rem; font-size: 1.1rem;">
                Configuration, Testing & Debugging Center
            </p>
        </div>
        """)
        
        with gr.Tabs():
            # Main Chat Tab
            with gr.TabItem("💬 Chat Interface"):
                with gr.Row():
                    with gr.Column(scale=3):
                        chatbot = gr.Chatbot(
                            height=600,
                            show_label=False,
                            avatar_images=["🧑‍💼", "🤖"]
                        )
                        
                        with gr.Row():
                            msg = gr.Textbox(
                                placeholder="Ask anything...",
                                show_label=False,
                                scale=4
                            )
                            send_btn = gr.Button("Send →", scale=1)
                        
                        clear_btn = gr.Button("🗑️ Clear Chat")
                    
                    with gr.Column(scale=1):
                        gr.Markdown("### 📚 Documents")
                        file_upload = gr.File(
                            label="Upload Files",
                            file_count="multiple",
                            file_types=[".txt", ".pdf", ".md", ".docx"]
                        )
                        upload_btn = gr.Button("📤 Load Documents", variant="primary")
                        upload_status = gr.Textbox(label="Status", interactive=False)
                        file_details = gr.Textbox(label="Details", lines=5, interactive=False)
            
            # Configuration Tab
            with gr.TabItem("⚙️ Configuration"):
                with gr.Row():
                    with gr.Column():
                        gr.Markdown("### Search Settings")
                        threshold_slider = gr.Slider(
                            minimum=0.0,
                            maximum=1.0,
                            value=rag.config['similarity_threshold'],
                            step=0.01,
                            label="Similarity Threshold",
                            info="Lower = more results, Higher = stricter matching"
                        )
                        
                        top_k_slider = gr.Slider(
                            minimum=1,
                            maximum=20,
                            value=rag.config['top_k_results'],
                            step=1,
                            label="Top K Results",
                            info="Number of chunks to retrieve"
                        )
                        
                        gr.Markdown("### Chunking Settings")
                        chunk_size_slider = gr.Slider(
                            minimum=100,
                            maximum=5000,
                            value=rag.config['chunk_size'],
                            step=100,
                            label="Chunk Size (chars)",
                            info="Size of text chunks"
                        )
                        
                        overlap_slider = gr.Slider(
                            minimum=0,
                            maximum=1000,
                            value=rag.config['chunk_overlap'],
                            step=50,
                            label="Chunk Overlap (chars)",
                            info="Overlap between chunks"
                        )
                    
                    with gr.Column():
                        gr.Markdown("### Features")
                        resume_mode = gr.Checkbox(
                            value=rag.config['enable_resume_mode'],
                            label="Resume Mode",
                            info="Special handling for resume/CV files"
                        )
                        
                        debug_mode = gr.Checkbox(
                            value=rag.config['debug_mode'],
                            label="Debug Mode",
                            info="Show debug info in responses"
                        )
                        
                        gr.Markdown("### Context Settings")
                        context_length = gr.Slider(
                            minimum=500,
                            maximum=8000,
                            value=rag.config['max_context_length'],
                            step=500,
                            label="Max Context Length",
                            info="Maximum context sent to LLM"
                        )
                        
                        model_dropdown = gr.Dropdown(
                            choices=["llama3.2", "mistral", "phi3", "deepseek-coder"],
                            value=rag.config['model'],
                            label="Model"
                        )
                
                with gr.Row():
                    save_config_btn = gr.Button("💾 Save Configuration", variant="primary")
                    load_config_btn = gr.Button("📂 Load Configuration")
                    config_status = gr.Textbox(label="Config Status", interactive=False)
            
            # Testing Tab
            with gr.TabItem("🧪 Query Testing"):
                gr.Markdown("### Test Semantic Search")
                
                test_query = gr.Textbox(
                    label="Test Query",
                    placeholder="Enter a query to test semantic search..."
                )
                test_btn = gr.Button("🔍 Run Test", variant="primary")
                
                test_output = gr.Textbox(
                    label="Test Results",
                    lines=20,
                    interactive=False,
                    elem_classes="debug-output"
                )
                
                gr.Markdown("### Quick Tests")
                with gr.Row():
                    gr.Examples(
                        examples=[
                            "experience",
                            "skills",
                            "education",
                            "create a bio",
                            "summarize this",
                            "what are the main points"
                        ],
                        inputs=test_query
                    )
            
            # Debug Tab
            with gr.TabItem("🐛 Debug & Stats"):
                with gr.Row():
                    with gr.Column():
                        gr.Markdown("### System Status")
                        refresh_btn = gr.Button("🔄 Refresh Status")
                        system_status = gr.Textbox(
                            label="Status",
                            lines=20,
                            interactive=False
                        )
                    
                    with gr.Column():
                        gr.Markdown("### Actions")
                        export_btn = gr.Button("📥 Export Chunks")
                        export_status = gr.Textbox(label="Export Status", interactive=False)
                        
                        gr.Markdown("### Chunk Viewer")
                        chunk_index = gr.Number(label="Chunk Index", value=0, precision=0)
                        view_chunk_btn = gr.Button("👁️ View Chunk")
                        chunk_viewer = gr.Textbox(
                            label="Chunk Content",
                            lines=10,
                            interactive=False
                        )
        
        # Event handlers
        def update_config(threshold, top_k, chunk_size, overlap, resume, debug, context, model):
            rag.update_config(
                similarity_threshold=threshold,
                top_k_results=top_k,
                chunk_size=chunk_size,
                chunk_overlap=overlap,
                enable_resume_mode=resume,
                debug_mode=debug,
                max_context_length=context,
                model=model
            )
            return "✅ Configuration updated"
        
        def view_chunk(index):
            try:
                if 0 <= index < len(rag.chunks):
                    chunk = rag.chunks[int(index)]
                    metadata = rag.chunk_metadata[int(index)]
                    return f"Source: {metadata['source']}\nType: {metadata['type']}\nSize: {metadata['size']} chars\n\n{chunk}"
                return "Invalid chunk index"
            except:
                return "No chunks loaded"
        
        # Connect configuration controls
        for control in [threshold_slider, top_k_slider, chunk_size_slider, 
                       overlap_slider, resume_mode, debug_mode, context_length, model_dropdown]:
            control.change(
                update_config,
                inputs=[threshold_slider, top_k_slider, chunk_size_slider, 
                       overlap_slider, resume_mode, debug_mode, context_length, model_dropdown],
                outputs=config_status
            )
        
        # File upload
        def load_and_update_stats(files):
            status, details, stats = rag.load_documents(files)
            return status, details, rag.get_system_status()
        
        upload_btn.click(
            load_and_update_stats,
            inputs=[file_upload],
            outputs=[upload_status, file_details, system_status]
        )
        
        # Chat
        msg.submit(lambda m, h: ("", h) if not m.strip() else ("", list(rag.stream_response(m, h))[-1]), 
                  [msg, chatbot], [msg, chatbot])
        send_btn.click(lambda m, h: ("", h) if not m.strip() else ("", list(rag.stream_response(m, h))[-1]), 
                      [msg, chatbot], [msg, chatbot])
        clear_btn.click(lambda: ([], ""), outputs=[chatbot, msg])
        
        # Testing
        test_btn.click(rag.test_query, inputs=test_query, outputs=test_output)
        
        # Debug
        refresh_btn.click(rag.get_system_status, outputs=system_status)
        export_btn.click(rag.export_chunks, outputs=export_status)
        view_chunk_btn.click(view_chunk, inputs=chunk_index, outputs=chunk_viewer)
        
        # Config save/load
        save_config_btn.click(rag.save_config, outputs=config_status)
        load_config_btn.click(rag.load_config, outputs=config_status)
    
    return app

if __name__ == "__main__":
    print("\n" + "="*60)
    print("🎛️ RAG ADMIN DASHBOARD")
    print("="*60)
    print("Features:")
    print("  • Adjustable similarity threshold")
    print("  • Configurable chunk sizes")
    print("  • Query testing with scores")
    print("  • Debug information")
    print("  • Configuration save/load")
    print("  • Chunk viewer")
    print("  • System status monitor")
    print("="*60 + "\n")
    
    app = create_admin_dashboard()
    app.launch(server_name="0.0.0.0", server_port=7862)